from __future__ import annotations

import logging
import hashlib
import re
from datetime import date, datetime
from io import BytesIO
from typing import Iterable

import requests
from bs4 import BeautifulSoup
from docx import Document

from ..models import SourceDocument, SourceRecord
from ..normalization import (
    district_from_text,
    neighborhood_from_text,
    normalize_parcels,
    plan_scales_from_text,
    stable_hash,
)
from .http import http_session

logger = logging.getLogger(__name__)

COUNCIL_URL = "https://www.ankara.bel.tr/meclis/kararlar"
TIMEOUT = 90
RELEVANT_TERMS = (
    "imar",
    "parsel",
    "nazım",
    "uygulama plan",
    "plan not",
    "imar ve bayındırlık",
)


def _parse_date(value: str) -> str:
    return datetime.strptime(value.strip(), "%d.%m.%Y").date().isoformat()


def _relevant(summary: str) -> bool:
    folded = summary.casefold()
    return any(term.casefold() in folded for term in RELEVANT_TERMS)


def parse_council_html(
    html: str, *, relevant_only: bool = True
) -> list[SourceRecord]:
    soup = BeautifulSoup(html, "html.parser")
    records: list[SourceRecord] = []
    for row in soup.select("table tbody tr"):
        cells = row.find_all("td")
        if len(cells) < 3:
            continue
        raw_date = cells[0].get_text(" ", strip=True)
        decision_number = cells[1].get_text(" ", strip=True)
        summary = cells[2].get_text(" ", strip=True)
        if relevant_only and not _relevant(summary):
            continue
        link = row.find("a", href=True)
        document_url = link["href"].strip() if link else ""
        if not raw_date or not decision_number or not summary:
            continue
        event_date = _parse_date(raw_date)
        source_id = f"council:{event_date}:{decision_number}"
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if document_url.casefold().endswith(".docx")
            else "application/pdf"
        )
        record = SourceRecord(
            source_id=source_id,
            source_kind="council",
            source_type="ABB_COUNCIL",
            title=summary,
            event_date=event_date,
            district=district_from_text(summary),
            neighborhood=neighborhood_from_text(summary),
            parcels=normalize_parcels(summary),
            plan_scales=plan_scales_from_text(summary),
            decision_number=decision_number,
            documents={
                "primary": SourceDocument(
                    url=document_url, media_type=media_type
                )
            },
            raw={"summary": summary},
        )
        record.snapshot_hash = stable_hash(
            {
                "source_id": source_id,
                "summary": summary,
                "document_url": document_url,
            }
        )
        records.append(record)
    return records


def fetch_council_records(
    *,
    from_date: date,
    max_pages: int = 200,
    session: requests.Session | None = None,
) -> tuple[list[SourceRecord], list[str]]:
    http = session or http_session()
    records: list[SourceRecord] = []
    errors: list[str] = []
    seen: set[str] = set()
    for page in range(1, max_pages + 1):
        try:
            response = http.get(
                COUNCIL_URL,
                params={"page": page},
                timeout=TIMEOUT,
                headers={"Referer": "https://www.ankara.bel.tr/"},
            )
            response.raise_for_status()
            page_records = parse_council_html(response.text, relevant_only=False)
        except (requests.RequestException, ValueError) as exc:
            errors.append(f"page={page}: {exc}")
            logger.warning("Council page failed: %s", errors[-1])
            break
        if not page_records:
            break
        oldest = min(date.fromisoformat(item.event_date) for item in page_records)
        for record in page_records:
            if (
                date.fromisoformat(record.event_date) >= from_date
                and _relevant(record.title)
                and record.source_id not in seen
            ):
                seen.add(record.source_id)
                records.append(record)
        if oldest < from_date:
            break
    logger.info("Council: %d relevant record(s)", len(records))
    return records, errors


def extract_docx_text(content: bytes) -> tuple[str, list[dict[str, object]]]:
    """Extract paragraphs/tables while preserving explicit strikethrough runs."""
    document = Document(BytesIO(content))
    blocks: list[str] = []
    evidence: list[dict[str, object]] = []
    paragraph_index = 0

    def render_paragraph(paragraph) -> str:
        nonlocal paragraph_index
        paragraph_index += 1
        pieces: list[str] = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            pieces.append(f"~~{text}~~" if run.font.strike else text)
        rendered = "".join(pieces).strip()
        if rendered:
            evidence.append(
                {"paragraph": paragraph_index, "text": rendered[:1000]}
            )
        return rendered

    for paragraph in document.paragraphs:
        rendered = render_paragraph(paragraph)
        if rendered:
            blocks.append(rendered)
    for table in document.tables:
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_text = " ".join(
                    value
                    for paragraph in cell.paragraphs
                    if (value := render_paragraph(paragraph))
                )
                if cell_text:
                    cells.append(cell_text)
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks), evidence


def download_document(
    document: SourceDocument, session: requests.Session | None = None
) -> bytes:
    http = session or http_session()
    response = http.get(document.url, timeout=TIMEOUT)
    response.raise_for_status()
    content = response.content
    document.size_bytes = len(content)
    document.sha256 = hashlib.sha256(content).hexdigest()
    return content


def iter_recent_records(
    records: Iterable[SourceRecord], cutoff: date
) -> Iterable[SourceRecord]:
    return (
        record
        for record in records
        if date.fromisoformat(record.event_date) >= cutoff
    )
